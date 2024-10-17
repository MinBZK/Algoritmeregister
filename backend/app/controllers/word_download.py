import datetime
import io
from typing import Tuple
from docx import Document, document
from fastapi.responses import StreamingResponse
from fastapi import HTTPException, status
from sqlalchemy.orm import Session
from app import schemas
from app.schemas.config.types import SchemaProperty
from app.config.layouts.types import LayoutJson
from app.schemas.misc import Language
from app.util.config_load import get_ttl_hash, collect_structure_data
from app.repositories import AlgoritmeVersionRepository


def get_org_data(db: Session, org_name: str) -> list[schemas.AlgoritmeVersionDownload]:
    algoritme_version_repository = AlgoritmeVersionRepository(db)
    algorithms = algoritme_version_repository.get_latest_by_org_by_lang(
        org_name, Language.NLD
    )
    return [schemas.AlgoritmeVersionDownload(**row.dict()) for row in algorithms]


def get_algo_data(db: Session, lars: str) -> schemas.AlgoritmeVersionDownload | None:
    algoritme_version_repository = AlgoritmeVersionRepository(db)
    algorithm = algoritme_version_repository.get_latest_by_lars_by_lang(
        lars, Language.NLD
    )
    if algorithm:
        return schemas.AlgoritmeVersionDownload(**algorithm.dict())


def insert_rich_text_field(doc, field: str):
    from bs4 import BeautifulSoup, Tag

    soup = BeautifulSoup(field, "html.parser")
    current_key_type = None

    for tag in soup.descendants:
        if isinstance(tag, Tag):
            if tag.name in ["div", "p"]:
                if tag.find("p"):
                    continue
                else:
                    text = tag.get_text(strip=True)
                    if text:
                        doc.add_paragraph(text)
            elif tag.name == "ul":
                current_key_type = "List Bullet"
            elif tag.name == "ol":
                current_key_type = "List Number"
            elif tag.name == "li":
                text = tag.get_text(strip=True)
                if current_key_type:
                    doc.add_paragraph(text, style=current_key_type)
            elif tag.name == "br":
                doc.add_paragraph("")

    return doc


def insert_list_field(doc: document.Document, field: list[str] | list[dict]):
    if type(field[0]) == str:
        if len(field) > 1:
            for bullet in field:
                doc.add_paragraph(str(bullet), style="List Bullet")
        else:
            doc.add_paragraph(field[0])
    elif type(field[0]) == dict:
        if len(field) > 1:
            for bullet in field:
                stringified_dict = ", ".join(filter(None, dict(bullet).values()))  # type: ignore
                doc.add_paragraph(stringified_dict, style="List Bullet")
        else:
            stringified_dict = ", ".join(filter(None, field[0].values()))
            doc.add_paragraph(stringified_dict)
    return doc


def insert_styled_field(
    doc: document.Document,
    field: str | list[str] | list[dict],
    *,
    is_rich_text: bool = False,
):
    if type(field) == str:
        if is_rich_text:
            return insert_rich_text_field(doc, field)
        doc.add_paragraph(field)
        return doc
    elif type(field) == list:
        return insert_list_field(doc, field)
    else:
        return doc


def build_description(
    doc: document.Document,
    algorithm: schemas.AlgoritmeVersionDownload,
    layout: LayoutJson,
    schema: dict[str, SchemaProperty],
):
    doc.add_heading(f"{algorithm.name} ({algorithm.lars})", 1)
    p = doc.add_paragraph("Publicatiestandaard: ")
    p.add_run(f"{algorithm.standard_version}").bold = True

    for tab_group in layout.tabsGrouping:
        doc.add_heading(f"{tab_group.label}", 2)

        field_list = tab_group.rows
        for field in field_list:
            doc.add_heading(f"{schema[field].title}", 3)

            field_value = getattr(algorithm, field)
            is_rich_text = schema[field].allowed_html_tags is not None
            if field_value is None:
                field_value = ""
                is_rich_text = False

            doc = insert_styled_field(doc, field_value, is_rich_text=is_rich_text)

    return doc


def get_all_algorithm_doc(db: Session, org_name: str):
    data = get_org_data(db, org_name)
    if len(data) == 0:
        return None, None

    standards, layouts = collect_structure_data(get_ttl_hash())
    organisation = data[0].organization
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d")

    doc = Document()
    doc.add_heading("Algoritmebeschrijvingen", 0)
    doc.add_paragraph(f"Algoritmes van {organisation}", style="Subtitle")
    doc.add_paragraph(
        f"""Dit document bevat algoritmebeschrijvingen van {organisation}, uitgedraaid op {timestamp}. \
Gebruik het om algoritmebeschrijvingen intern te reviewen en wijzigingen te accorderen voor publicatie \
op het algoritmeregister van de Nederlandse overheid. Heeft u vragen of heeft u hulp nodig? Neem via mail \
contact op met: algoritmeregister@minbzk.nl."""
    )

    for algorithm in data:
        doc.add_page_break()
        if not algorithm.standard_version:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT, detail="STANDARD_NOT_KNOWN"
            )
        layout = layouts[algorithm.standard_version]
        schema = standards[algorithm.standard_version]

        doc = build_description(doc, algorithm, layout, schema)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    filename = f"Algoritmebeschrijvingen van {organisation} op {timestamp}.docx"
    return stream, filename


def get_one_algorithm_doc(db, lars: str) -> Tuple[io.BytesIO | None, str | None]:
    algorithm = get_algo_data(db, lars)
    if algorithm is None:
        return None, None

    standards, layouts = collect_structure_data(get_ttl_hash())
    organisation = algorithm.organization
    name = algorithm.name
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d")

    doc: document.Document = Document()
    doc.add_heading("Algoritmebeschrijving", 0)
    doc.add_paragraph(f"Algoritme: {name} van {organisation}", style="Subtitle")
    doc.add_paragraph(
        f"""Dit document bevat de algoritmebeschrijving: {name} van {organisation}, uitgedraaid op {timestamp}. \
Gebruik het om algoritmebeschrijvingen intern te reviewen en wijzigingen te accorderen voor publicatie \
op het algoritmeregister van de Nederlandse overheid. Heeft u vragen of heeft u hulp nodig? Neem via mail \
contact op met: algoritmeregister@minbzk.nl."""
    )

    if not algorithm.standard_version:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT, detail="STANDARD_NOT_KNOWN"
        )
    layout = layouts[algorithm.standard_version]
    schema = standards[algorithm.standard_version]

    doc = build_description(doc, algorithm, layout, schema)

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)

    filename = f"{name.encode('utf-8')} {timestamp}.docx"
    return stream, filename


def generate_word_download(
    db: Session, *, org_name: str | None = None, lars: str | None = None
):
    if lars:
        stream, filename = get_one_algorithm_doc(db, lars)
    elif org_name:
        stream, filename = get_all_algorithm_doc(db, org_name)
    else:
        raise ValueError("Please enter one of two identifiers: lars | org_name")

    if stream is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="No algorithm descriptions found. Unable to make document",
        )
    media_type = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response = StreamingResponse(io.BytesIO(stream.read()), media_type=media_type)
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


if __name__ == "__main__":
    collect_structure_data()
